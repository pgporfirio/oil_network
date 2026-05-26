"""Update v2.3's hard-coded TOC and List of Tables to reflect the spliced
new §5.4 (Genscape worked example).

The TOC entries in the docx are hard-coded paragraphs (not Word TOC fields).
After splicing the new section, we need to:
  - Insert a new TOC entry for §5.4 "Multi-source integration..."
  - Renumber the displaced entries: §5.4 → §5.5, §5.5 → §5.6
  - Shift the page numbers of all entries from §5.4 onward by the page delta
    introduced by the new section
  - Update Table 5.3 entry in List of Tables to Table 5.4, and insert a new
    Table 5.3 entry for the Genscape comparison

Reads page numbers from the freshly-rendered v2.3.pdf so the update is anchored
to the actual rendered pagination, not a guess.
"""
from __future__ import annotations
from pathlib import Path
import re
import sys

import fitz
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

DOCS = Path(__file__).resolve().parent.parent / "outputs" / "docs"
DOCX = DOCS / "Master_Thesis_Pedro_Porfirio_restructured_v2.3.docx"
PDF  = DOCS / "Master_Thesis_Pedro_Porfirio_restructured_v2.3.pdf"

# Front-matter offset: how many PDF pages precede the arabic-numbered body.
# Inferred from rendering: §1.1 lives on PDF p12 which is printed page 1.
FRONT_OFFSET = 11


def find_pages(pdf_path: Path):
    """Scan the PDF and return a dict mapping section markers to printed page
    numbers. Markers recognised:
      - Chapter 5 sections: '5.1 ...', '5.2 ...', etc.
      - Chapter 6 sections: '6.1 ...', etc., plus the chapter heading
      - 'References'
      - 'Annex A', 'Annex B', 'Annex C' and their subsections
      - The new §5.4 heading (matched by substring "Multi-source integration")
    """
    doc = fitz.open(str(pdf_path))
    pages = {}

    def claim(key, page):
        # Take the earliest page that holds the heading.
        if key not in pages:
            pages[key] = page - FRONT_OFFSET

    # Patterns we care about (compiled once)
    chapter_pat   = re.compile(r"^([5-6])\s+[A-Z]")
    section_pat   = re.compile(r"^([5-6]\.\d+)\s+[A-Z]")
    annex_chap    = re.compile(r"^Annex\s+([A-C])")
    annex_section = re.compile(r"^([A-C]\.\d+)\s+[A-Z]")
    multi_source  = re.compile(r"Multi-source integration")

    for i in range(FRONT_OFFSET, len(doc)):
        text = doc[i].get_text()
        for line in text.split("\n"):
            s = line.strip()
            if not s:
                continue
            if multi_source.search(s):
                claim("5.4_new", i + 1)
            m = section_pat.match(s)
            if m:
                claim(m.group(1), i + 1)
                continue
            m = chapter_pat.match(s)
            if m:
                claim(f"chapter_{m.group(1)}", i + 1)
                continue
            m = annex_section.match(s)
            if m:
                claim(m.group(1), i + 1)
                continue
            m = annex_chap.match(s)
            if m:
                claim(f"annex_{m.group(1)}", i + 1)
                continue
            if s == "References":
                claim("References", i + 1)
    doc.close()
    return pages


# Mapping: TOC entry text prefix → key in `pages` dict.
# For the new §5.4, we insert a fresh entry rather than rewriting a label.
TOC_UPDATES = [
    # (prefix_to_match, new_text_first_run, page_key)
    ("5.1",                  "5.1   Setup", "5.1"),
    ("5.2",                  "5.2   Steady-state comparison: the PADD 5 in-transit anomaly", "5.2"),
    ("5.3",                  "5.3   Stress-test comparison: Hurricane Harvey, August 2017", "5.3"),
    # ("5.4", ...) — will be the displaced one, renumbered to 5.5 below
    # ("5.5", ...) — will be the displaced one, renumbered to 5.6 below
    ("6   Conclusion",       "6   Conclusion and future work", "chapter_6"),
    ("6.1",                  "6.1   Contribution", "6.1"),
    ("6.2",                  "6.2   Limitations", "6.2"),
    ("6.3",                  "6.3   Future work", "6.3"),
    ("6.4",                  "6.4   Closing remarks", "6.4"),
    ("References",           "References", "References"),
    ("Annex A",              None, "annex_A"),  # text stays, just page shifts
    ("A.1",                  "A.1   Schema", "A.1"),
    ("A.2",                  "A.2   The starter U.S. crude oil graph", "A.2"),
    ("A.3",                  "A.3   The resolver", "A.3"),
    ("A.4",                  "A.4   Consistency properties demonstrated on the live system", "A.4"),
    ("Annex B",              None, "annex_B"),
    ("B.1",                  "B.1   Variable types", "B.1"),
    ("B.2",                  "B.2   Authoritative bindings by subsystem", "B.2"),
    ("B.3",                  "B.3   Latent variables by location", "B.3"),
    ("Annex C",              None, "annex_C"),
    ("C.1",                  "C.1   Node embeddings and message passing", "C.1"),
    ("C.2",                  "C.2   Graph convolutional networks (GCN)", "C.2"),
    ("C.3",                  "C.3   Temporal extensions and hybrid architectures", "C.3"),
    ("C.4",                  "C.4   Why the framework is GNN-ready", "C.4"),
]


def set_toc_runs(p, label_text, page_str):
    """Set the two runs of a TOC paragraph: heading text + tab+page number.
    Preserves boldness of existing runs."""
    runs = p.runs
    if not runs:
        return
    was_bold = runs[0].bold
    # Use the first run for label, second for page number; drop extras.
    runs[0].text = label_text
    if len(runs) >= 2:
        runs[1].text = f"\t{page_str}"
    # Trim extra runs that may exist
    if len(runs) > 2:
        for r in runs[2:]:
            r.text = ""
    # Re-apply bold (some chapter-level entries are bold)
    if was_bold:
        for r in p.runs[:2]:
            r.bold = True


def update_toc(d, pages):
    """Walk the TOC region and patch each known entry."""
    # Locate the §5.4 entry (currently "5.4   The same pattern...") and
    # insert a new TOC entry for the new §5.4 immediately BEFORE it.
    toc_paras = d.paragraphs
    # Index map by prefix match (heading runs[0].text starts-with)
    by_prefix = {}
    for i, p in enumerate(toc_paras):
        # Only paragraphs that have a tab+number second-run look like TOC entries.
        if len(p.runs) < 2:
            continue
        if not p.runs[1].text.startswith("\t"):
            continue
        first = p.runs[0].text.strip()
        by_prefix[first[:35]] = (i, p)

    # 1. Insert new TOC entry for §5.4 Multi-source integration.
    target_label = "5.4   The same pattern"
    target_idx_para = None
    for prefix, (i, p) in by_prefix.items():
        if prefix.startswith(target_label[:20]):
            target_idx_para = p
            break
    if target_idx_para is None:
        raise LookupError("Could not find TOC entry for old §5.4")

    new_p = OxmlElement("w:p")
    # Run 1: label
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.text = ("5.4   Multi-source integration: layering pipeline-flow data on "
               "top of EIA aggregates")
    t1.set(qn("xml:space"), "preserve")
    r1.append(t1)
    new_p.append(r1)
    # Run 2: tab + page
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.text = f"\t{pages['5.4_new']}"
    t2.set(qn("xml:space"), "preserve")
    r2.append(t2)
    new_p.append(r2)
    target_idx_para._element.addprevious(new_p)

    # 2. Rewrite the displaced TOC entries: §5.4 → §5.5, §5.5 → §5.6
    # (re-scan because we inserted)
    for i, p in enumerate(d.paragraphs):
        if len(p.runs) < 2 or not p.runs[1].text.startswith("\t"):
            continue
        first = p.runs[0].text.strip()
        if first.startswith("5.4   The same pattern"):
            # This is the OLD §5.4 — now §5.5
            page = pages.get("5.5")  # 5.5 in the pages dict is the renumbered (the OLD §5.4 body content)
            # But our PDF scan picked up "5.5" → printed page of "5.5 The same pattern"
            # which is what we want here
            set_toc_runs(p, "5.5   The same pattern on other named events", str(page) if page else "")
        elif first.startswith("5.5   What the comparison"):
            page = pages.get("5.6")
            set_toc_runs(p, "5.6   What the comparison establishes", str(page) if page else "")

    # 3. Update existing TOC entries from §5.1 onward where pages shifted.
    for prefix, new_label, page_key in TOC_UPDATES:
        target = None
        for i, p in enumerate(d.paragraphs):
            if len(p.runs) < 2 or not p.runs[1].text.startswith("\t"):
                continue
            first = p.runs[0].text.strip()
            if first.startswith(prefix):
                target = p
                break
        if target is None:
            print(f"  WARN: no TOC entry matching {prefix!r}")
            continue
        page = pages.get(page_key)
        if page is None:
            print(f"  WARN: no page number found for {page_key!r}")
            continue
        if new_label is None:
            # Only update page number, leave text alone
            target.runs[1].text = f"\t{page}"
        else:
            set_toc_runs(target, new_label, str(page))


def update_list_of_tables(d, pages):
    """Update List of Tables: insert new Table 5.3 entry (Genscape), renumber
    old Table 5.3 to Table 5.4."""
    # Find the "Table 5.3" entry in the LoT region.
    for i, p in enumerate(d.paragraphs):
        if len(p.runs) < 2 or not p.runs[1].text.startswith("\t"):
            continue
        first = p.runs[0].text.strip()
        if first.startswith("Table 5.3"):
            # Capture for renumbering, then insert a new Table 5.3 entry before it.
            old_table_53_p = p
            # New Table 5.3 entry (Genscape) — page is wherever §5.4 starts
            new_p = OxmlElement("w:p")
            r1 = OxmlElement("w:r")
            t1 = OxmlElement("w:t")
            t1.text = ("Table 5.3   Layering Genscape pipeline-flow data on top of EIA monthly "
                       "aggregates for PADD 2")
            t1.set(qn("xml:space"), "preserve")
            r1.append(t1)
            new_p.append(r1)
            r2 = OxmlElement("w:r")
            t2 = OxmlElement("w:t")
            # Table 5.3 lives on the same page as the new §5.4 or one page later.
            # Take the §5.4 start page.
            t2.text = f"\t{pages['5.4_new']}"
            t2.set(qn("xml:space"), "preserve")
            r2.append(t2)
            new_p.append(r2)
            old_table_53_p._element.addprevious(new_p)
            # Rewrite the displaced entry as Table 5.4 with new page
            new_table_54_page = pages.get("5.5")  # same section as the now-Table-5.4
            set_toc_runs(old_table_53_p,
                         "Table 5.4   Absolute value of B at the USA aggregate around named disruption events",
                         str(new_table_54_page) if new_table_54_page else "")
            return
    print("  WARN: could not find Table 5.3 entry in List of Tables")


def main():
    if not PDF.exists():
        sys.exit(f"v2.3 PDF not found: {PDF}")
    pages = find_pages(PDF)
    print("Detected printed pages:")
    for k in sorted(pages):
        print(f"  {k:20s} -> printed page {pages[k]}")

    d = docx.Document(str(DOCX))
    update_toc(d, pages)
    update_list_of_tables(d, pages)
    d.save(str(DOCX))  # overwrite v2.3 in-place
    print(f"\nUpdated TOC + LoT in {DOCX.name}")


if __name__ == "__main__":
    main()
