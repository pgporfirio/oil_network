"""
Rewrite Chapter 5 of v2.5 thesis → v2.6.

Changes:
1. §5.2 retitled "the PADD 5 in-transit anomaly" → "the Jones Act in-transit anomaly"
2. §5.2 prose fully rewritten (9 new paragraphs from drafted text)
3. Table 5.1 cells updated to reflect both-methods-label-correctly framing
4. New Figure 5.1 inserted (per-grade extension diagram)
5. §5.4 Genscape table caption + in-text reference: "Table 5.4" → "Table 5.3"
6. §5.5 Harvey magnitude clarifier inserted (Harvey is sub-median at USA agg
   because PADD 3 disruption partly cancels with offsetting national flows).
"""
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path(r'C:\Users\PedroPorfirio\OneDrive - Jabuticaba\Oil Network Project\Stage1\outputs\docs\Master_Thesis_Pedro_Porfirio_restructured_v2.5.docx')
DST = Path(r'C:\Users\PedroPorfirio\OneDrive - Jabuticaba\Oil Network Project\Stage1\outputs\docs\Master_Thesis_Pedro_Porfirio_restructured_v2.6.docx')
FIG = Path(r'C:\Users\PedroPorfirio\OneDrive - Jabuticaba\Oil Network Project\Stage1\outputs\docs\figures\fig_5_1_jones_act_corridor.png')

# =================== new prose for §5.2 ===================
P1 = (
    "EIA's stock identity for PADD 5 — the West Coast PADD, comprising California, "
    "Washington, Oregon, Alaska, Nevada, Arizona, and Hawaii — includes commercial "
    "stocks at tank farms, refinery stocks, the Strategic Petroleum Reserve allocation "
    "to PADD 5 (zero in practice, as SPR sites are on the Gulf Coast), and an "
    "in-transit-by-water allocation. The published PADD 5 stock total exceeds the sum "
    "of named PADD 5 hubs and refinery stocks by approximately six million barrels in "
    "every month of the modelled period. The residual is not an artefact. PADD 5 is "
    "structurally a net importer with limited pipeline connectivity to the rest of "
    "the United States, and a meaningful share of the crude consumed at West Coast "
    "refineries arrives by tanker from PADD 3 under the Jones Act — the cabotage law "
    "that requires U.S.-flagged, U.S.-built, and U.S.-crewed vessels for domestic "
    "point-to-point trade. The voyage from Houston or Beaumont to Los Angeles or "
    "San Francisco takes approximately ten days, so at any given month-end a "
    "non-trivial volume of PADD-3-origin crude is in transit between the two PADDs. "
    "EIA reports this volume separately in its in-transit-by-water series for PADD 5, "
    "and allocates it to PADD 5 as destination-based ending stocks."
)

P2 = (
    "A competent spreadsheet analyst handles this correctly. The EIA in-transit-by-water "
    "series is published at PADD level, so the workbook ingests it as a separate row "
    "and labels it — most commonly “in-transit” or “Jones Act in-transit "
    "(PADD 3 origin)” — and the PADD 5 stock identity closes exactly. The "
    "framework handles the same volume identically at the level of the published stock "
    "figure: the in-transit series binds to the inventory variable of a named asset, "
    "inter_padd_3_to_5_agg, and the PADD 5 partition closes. Both methods agree on the "
    "total, both name the six million barrels, both close. The labelling itself is not "
    "what distinguishes the two methods, as Table 5.1 lays out."
)

# Updated Table 5.1 caption
T51_CAPTION = (
    "Table 5.1. PADD 5 stock identity, representative month, with both methods binding "
    "to the EIA in-transit-by-water series. The published total closes in both columns; "
    "the methods agree on the number and on the label. The structural difference "
    "between them is visible in Figure 5.1, not at this level."
)

P4 = (
    "The methods diverge in where the labelled number lives. In the spreadsheet, the "
    "row is a flat cell with no neighbours: the analyst can read its value at any "
    "month-end, but there is no associated outflow from PADD 3, no inflow to PADD 5, "
    "no capacity ceiling, no mass-balance link. If a downstream consumer asks what "
    "the Jones Act northbound flow was in August 2017 — distinct from the in-transit "
    "stock at month-end — the workbook has no row to point at. The label is a name, "
    "not a position in the graph. Different workbooks label the same row differently, "
    "and two analysts working from the same EIA data can reach different downstream "
    "allocations depending on how they tie the in-transit row back to PADD 3 exports "
    "and PADD 5 imports. The label is correct; the interpretation is ad hoc."
)

P5 = (
    "In the framework, inter_padd_3_to_5_agg is a directed physical asset connecting "
    "the Gulf Coast complex to the West Coast PADD. Its inventory is the line-fill on "
    "the corridor — the six million barrels in transit. Its outflow variable on the "
    "PADD-3 side is the monthly volume leaving the Gulf bound for PADD 5; its inflow "
    "variable on the PADD-5 side is the volume arriving. Mass balance ties them: "
    "ΔS_corridor = F_in_from_PADD3 − F_out_to_PADD5 + B, with B carrying the "
    "irreducible measurement residual. The corridor has a tanker-fleet capacity "
    "ceiling, so its monthly throughput sits inside an explicit upper bound. The same "
    "six million barrels sit at the same place in the balance, but they now occupy a "
    "node in a graph rather than a row in a flat table. Figure 5.1 shows the same six "
    "million barrels in both representations side by side, with a stacked view below "
    "of what each method does when a grade dimension is added."
)

FIG_CAPTION = (
    "Figure 5.1. The same six million barrels in each method, with the per-grade "
    "extension shown below. Spreadsheet (left): a row in a flat table; adding a grade "
    "dimension duplicates every column and proliferates new residuals. Framework "
    "(right): a node with neighbours, capacity, per-direction flow variables, and "
    "mass balance; adding a grade dimension widens the variables on the same nodes "
    "and edges, leaving the topology unchanged."
)

P6 = (
    "The structural placement matters in two distinct ways. The first is short-range: "
    "as soon as a finer-resolution source becomes available, the framework absorbs it "
    "without restructuring. Kpler tracks Jones Act vessel movements through AIS feeds, "
    "and a future subscription to that data binds per-vessel cargo and ETA series to "
    "the same corridor node, replacing the EIA in-transit aggregate with vessel-level "
    "detail. Nothing else in the graph changes — the PADD 3 outflow variable, the "
    "PADD 5 inflow variable, the partition formulas, the mass-balance equations all "
    "carry on as before. The spreadsheet's equivalent is structural: vessel-level "
    "data has no row to bind to in the existing layout, so the workbook must grow a "
    "new column structure indexed by vessel, and the existing “in-transit” "
    "row becomes a sum over those columns whose definition the analyst must maintain."
)

P7 = (
    "The second way the placement matters is longer-range, and it tests the framework "
    "against a question commercial desks ask routinely: what grades are in transit, "
    "and how do they propagate through the system? The Jones Act northbound is not a "
    "single grade. WTI from Houston, Bakken transhipped through Pacific terminals, "
    "ANS heavy crude from Valdez, and various medium-sour cargoes from refining-region "
    "surplus all move through the corridor in different months, with grade-specific "
    "implications for what PADD 5 refiners can process and at what crack spreads. "
    "Answering the question requires the balance to be tracked per (asset, grade), "
    "not just per asset. The implementation in this thesis is single-grade; per-grade "
    "decomposition is the natural next data-layer extension, and the cost asymmetry "
    "between the two methods on that extension is the strongest test of the "
    "orthogonality principle in this chapter."
)

P8 = (
    "In the spreadsheet, per-grade decomposition is a rebuild. Every column in the "
    "existing balance must duplicate across the grade dimension. The "
    "“in-transit” row becomes a set of rows — one per grade — each requiring "
    "a sourcing rule that ties it to a specific PADD 3 export grade-band and a specific "
    "PADD 5 refinery slate. EIA's per-grade reporting is sparser than the aggregate: "
    "light sweet and medium sour are published by API-gravity band at PADD level, but "
    "the in-transit series is not decomposed by grade, so each grade's in-transit "
    "volume becomes a new residual that the analyst must allocate using assumptions "
    "outside the workbook. The same is true for every other line of the balance — "
    "production at every named basin must split into the grades that basin produces; "
    "refinery runs must split into grade-slate inputs and product-yield outputs; "
    "pipelines must carry grade-mix series; storage tanks must track "
    "per-(tank, grade) blends. The reconciliation rules that the original workbook "
    "encoded — every PADD-named-basin sum, every Jones Act in-transit allocation, "
    "every Cushing partition — must be redefined for each grade, and the residual "
    "cells multiply. The result is a new workbook by a new analyst, re-validated "
    "from scratch against everything the previous one did."
)

P9 = (
    "In the framework, the asset graph is unchanged. The 251 nodes, the 433 directed "
    "flow edges, the partition tree, and the capacity ceilings stay bit-identical. "
    "What changes is the data layer: the commodities table grows from one row to N, "
    "the variables table grows from 1,870 to approximately 1,870 × N, and the "
    "resolver runs over the wider variable set with no change to its dispatch rules. "
    "For each new variable: either a per-grade observed series binds where one is "
    "published, or the variable inherits the aggregate-level formula with the "
    "per-(node, grade) residual carried as B at the appropriate physical node. "
    "Partition closure now operates per (node, grade); v_aggregation_consistency "
    "flags divergence per (node, grade) automatically. Mass balance closes per "
    "(node, grade) at every physical asset. The refinery is the one place where the "
    "per-grade balance equation switches form — crude grades in differ from product "
    "grades out, because the refinery is the point of grade conversion — and that "
    "switch is handled either by a per-refinery yield formula (the extension flagged "
    "in §1.4) or, where a yield optimisation is required, by a constraint that "
    "the downstream LP consumer reads from the resolved table."
)

P10 = (
    "The orthogonality between the physical layer and the data layer is what makes "
    "this asymmetry possible. The spreadsheet conflates structure and data: every "
    "column is both what to track and what its value is, so widening the data "
    "dimension drags the structural layer along with it. The framework keeps the two "
    "apart — the asset graph encodes what to track, the variable assignments encode "
    "each variable's value for a given scenario — and the two layers evolve on "
    "independent timescales. A data-layer widening leaves the structural layer "
    "untouched."
)

P11 = (
    "The example matters because it shows what the orthogonality principle buys over "
    "the lifetime of the artefact, not at a single month-end. At the level of the "
    "published PADD 5 stock identity, the spreadsheet and the framework agree exactly. "
    "The framework's structural placement of the in-transit volume on a named "
    "physical asset pays off the first time a downstream consumer asks a question "
    "the EIA aggregate cannot answer — what was the per-vessel arrival schedule, "
    "what was the per-grade composition, what was the implied capacity utilisation — "
    "because the answer in the framework is a data binding on an existing node, "
    "while the answer in the spreadsheet is a workbook rewrite."
)


def make_para(text, *, italic=False, bold=False):
    """Create a w:p element with a single run carrying the text."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    if italic or bold:
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
        if italic:
            rPr.append(OxmlElement('w:i'))
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def find_and_replace_text(element, old, new):
    """Replace `old` with `new` in all w:t elements under `element`."""
    n = 0
    for t in element.findall('.//' + qn('w:t')):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            n += 1
    return n


def set_table_cell_text(cell, new_text):
    """Replace a table cell's text content entirely.
    Preserves the first paragraph + first run for styling; clears other runs."""
    paras = cell.findall(qn('w:p'))
    if not paras:
        # add a paragraph
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = new_text
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        cell.append(new_p)
        return
    # Replace text in first paragraph; remove additional paragraphs
    first_p = paras[0]
    # Remove all runs from first_p
    for r in first_p.findall(qn('w:r')):
        first_p.remove(r)
    # Add one new run with the text
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = new_text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    first_p.append(new_r)
    # Remove additional paragraphs
    for p in paras[1:]:
        cell.remove(p)


def main():
    doc = Document(str(SRC))
    body = doc.element.body
    elements = list(body)

    # =================== STEP 1: Update §5.2 heading title ===================
    heading_5_2 = elements[288]
    n = find_and_replace_text(heading_5_2,
                              'PADD 5 in-transit anomaly',
                              'Jones Act in-transit anomaly')
    print(f"[1] §5.2 heading: {n} replacement(s)")

    # =================== STEP 2: identify key §5.2 elements ===================
    prose_before_table = [elements[i] for i in (289, 290, 291, 292)]
    table_5_1 = elements[293]
    caption_5_1 = elements[294]
    prose_after_caption = [elements[i] for i in (295, 296, 297)]
    heading_5_3 = elements[298]

    # =================== STEP 3: Modify Table 5.1 cells ===================
    # Find the in-transit row by scanning for "In-transit" in cell text
    rows = table_5_1.findall(qn('w:tr'))
    print(f"[3] Table 5.1 has {len(rows)} rows")
    for row_idx, row in enumerate(rows):
        cells = row.findall(qn('w:tc'))
        cell_texts = [' '.join((t.text or '') for t in c.findall('.//' + qn('w:t'))).strip() for c in cells]
        # Show row content for debugging
        print(f"  row {row_idx}: {cell_texts}")

    # Map row updates:
    #   row containing "In-transit from PADD 3" → update label + both data cells
    #   row containing "Unallocated adjustment" → set both data cells to "0"
    #   row containing "Commercial tank farms" → drop the "~" prefix on spreadsheet cell
    #   row containing "Refinery stocks" → drop the "~" prefix on spreadsheet cell
    for row in rows:
        cells = row.findall(qn('w:tc'))
        if not cells:
            continue
        cell_text_0 = ' '.join((t.text or '') for t in cells[0].findall('.//' + qn('w:t'))).strip()

        if 'In-transit from PADD 3' in cell_text_0:
            set_table_cell_text(cells[0], 'In-transit by water from PADD 3')
            set_table_cell_text(cells[1], "5,900 (row labelled 'in-transit')")
            set_table_cell_text(cells[2], '5,900 (TS-bound to inter_padd_3_to_5_agg.inventory)')
            print("  → updated in-transit row")
        elif 'Unallocated adjustment' in cell_text_0:
            set_table_cell_text(cells[1], '0')
            set_table_cell_text(cells[2], '0')
            print("  → zeroed unallocated adjustment row")
        elif 'Commercial tank farms' in cell_text_0:
            # drop "~" from "~22,962"
            set_table_cell_text(cells[1], '22,962')
            print("  → cleaned commercial tank farms cell")
        elif 'Refinery stocks' in cell_text_0:
            set_table_cell_text(cells[1], '20,468')
            print("  → cleaned refinery stocks cell")

    # =================== STEP 4: Update caption 5.1 text ===================
    # Replace caption text entirely
    for t in caption_5_1.findall('.//' + qn('w:t')):
        caption_5_1.find('.//' + qn('w:r')).getparent().remove(caption_5_1.find('.//' + qn('w:r')))
        break  # only need to trigger; we'll rebuild below
    # remove all remaining runs
    for r in caption_5_1.findall('.//' + qn('w:r')):
        r.getparent().remove(r)
    # add a fresh run with new text
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = T51_CAPTION
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    caption_5_1.append(new_r)
    print(f"[4] Updated caption 5.1")

    # =================== STEP 5: Delete old prose paragraphs ===================
    for p in prose_before_table:
        body.remove(p)
    print(f"[5a] Deleted {len(prose_before_table)} prose paragraphs before table")
    for p in prose_after_caption:
        body.remove(p)
    print(f"[5b] Deleted {len(prose_after_caption)} prose paragraphs after caption")

    # =================== STEP 6: Insert new prose BEFORE table ===================
    # After heading_5_2: P1, P2 → then table follows
    prev = heading_5_2
    for text in (P1, P2):
        p = make_para(text)
        prev.addnext(p)
        prev = p
    print(f"[6] Inserted 2 prose paragraphs after §5.2 heading")

    # =================== STEP 7: Insert AFTER caption: P4, P5, Figure, FigCaption, P6..P11 ===================
    prev = caption_5_1

    # Insert P4, P5
    for text in (P4, P5):
        p = make_para(text)
        prev.addnext(p)
        prev = p

    # Insert Figure 5.1 paragraph
    # Use python-docx to append a centred paragraph with picture, then move
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_run = fig_para.add_run()
    fig_run.add_picture(str(FIG), width=Inches(6.0))
    # Move the figure paragraph element to immediately after `prev`
    fig_para_el = fig_para._element
    # python-docx places new paragraphs at the end; the element is in body BEFORE sectPr
    # We need to detach it then insert
    fig_para_el.getparent().remove(fig_para_el)
    prev.addnext(fig_para_el)
    prev = fig_para_el
    print(f"[7a] Inserted Figure 5.1")

    # Insert figure caption
    fig_cap = make_para(FIG_CAPTION)
    prev.addnext(fig_cap)
    prev = fig_cap

    # Insert P6..P11
    for text in (P6, P7, P8, P9, P10, P11):
        p = make_para(text)
        prev.addnext(p)
        prev = p
    print(f"[7b] Inserted Figure caption + 6 prose paragraphs after Figure 5.1")

    # =================== STEP 8: Renumber §5.4 Table 5.4 → Table 5.3 ===================
    # Find paragraphs containing "Table 5.4" and which are NOT the §5.5 named-events caption
    # The Genscape ones are the FIRST two occurrences (in-text ref + caption), the named-events
    # ones are the LAST two (in-text ref + caption).
    # By document order: §5.4 mentions first (idx 310 ref, 311 caption); §5.5 mentions last
    # (idx 316 ref, 319 caption — assuming original numbering).
    refreshed_elements = list(body)
    found_table_54_paragraphs = []
    for el in refreshed_elements:
        if el.tag != qn('w:p'):
            continue
        texts = ' '.join((t.text or '') for t in el.findall('.//' + qn('w:t')))
        if 'Table 5.4' in texts:
            found_table_54_paragraphs.append(el)
    print(f"[8] Found {len(found_table_54_paragraphs)} paragraphs containing 'Table 5.4'")
    # First two = §5.4 Genscape → rename to Table 5.3
    for i, el in enumerate(found_table_54_paragraphs[:2]):
        n = find_and_replace_text(el, 'Table 5.4', 'Table 5.3')
        print(f"  [8.{i+1}] Renamed Table 5.4 → Table 5.3 in §5.4 ({n} replacement(s))")
    # Last two = §5.5 named events → stays Table 5.4

    # =================== STEP 9: Insert Harvey clarifier into §5.5 prose ===================
    # Find the §5.5 paragraph that says "Second, around named events B rises notably above..."
    for el in refreshed_elements:
        if el.tag != qn('w:p'):
            continue
        texts = ' '.join((t.text or '') for t in el.findall('.//' + qn('w:t')))
        if 'B rises notably above the steady-state median' in texts:
            old_phrase = 'in the modelled period. The framework preserves this signal'
            new_phrase = (
                'in the modelled period. Hurricane Harvey at the USA aggregate is the '
                'partial exception: its national-level signal is partly cancelled by '
                'offsetting flows between PADD 3 and the rest of the country, while the '
                'per-PADD signal documented in §5.3 remains clear. The framework '
                'preserves this signal'
            )
            n = find_and_replace_text(el, old_phrase, new_phrase)
            print(f"[9] Inserted Harvey clarifier ({n} replacement(s))")
            break

    # =================== STEP 10: Save ===================
    doc.save(str(DST))
    print(f"\nSaved: {DST.name}")
    print(f"  source size: {SRC.stat().st_size:,} bytes")
    print(f"  output size: {DST.stat().st_size:,} bytes")


if __name__ == '__main__':
    main()
