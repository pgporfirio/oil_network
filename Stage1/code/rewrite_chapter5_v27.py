"""
Rewrite Chapter 5 of v2.6 → v2.7.

Reframes the chapter from steady-state numerical comparison to a four-scenario
change-management comparison across three methods: Excel spreadsheet, Pandas
notebook, framework.

Scenarios:
  §5.2 — Adding grade dimensions to basin production    (data-dimension widening)
  §5.3 — Refinery turnaround or unplanned outage         (resolution refinement, Harvey illustration)
  §5.4 — Receiving Genscape pipeline data                 (new data source + topology granularity)
  §5.5 — Finding inconsistencies                          (consistency-audit infrastructure)

Each scenario follows the same template: setup paragraph + three method narratives
+ one comparison table (4 columns × 6 rows). §5.6 aggregates with a cost-asymmetry table.

Figure 5.1 (Jones Act corridor + per-grade extension) stays in §5.2 with an
updated caption reframing it as the grade-widening illustration.
"""
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path(r'C:\Users\PedroPorfirio\OneDrive - Jabuticaba\Oil Network Project\Stage1\outputs\docs\Master_Thesis_Pedro_Porfirio_restructured_v2.6.docx')
DST = Path(r'C:\Users\PedroPorfirio\OneDrive - Jabuticaba\Oil Network Project\Stage1\outputs\docs\Master_Thesis_Pedro_Porfirio_restructured_v2.7.docx')

# =================== CHAPTER 5 CONTENT ===================

CHAPTER5_INTRO = (
    "The previous chapter set out the framework's design principles. This chapter "
    "compares the framework against the two dominant alternatives currently used "
    "to produce oil balances — the Excel spreadsheet workflow embedded in trading "
    "firms, consultancies, and government agencies, and the Python notebook workflow "
    "increasingly common on quantitative desks. The comparison is deliberately not "
    "about who produces a better single-month balance from the same data. On a fixed "
    "monthly EIA snapshot, all three methods, in the hands of a competent practitioner, "
    "arrive at approximately the same headline numbers. The interesting comparison is "
    "what happens when the world around the balance changes — when a new dimension "
    "is added to tracking, when a structural event breaks the assumptions the balance "
    "was built on, when a new data source becomes available, and when the analyst "
    "needs to verify that the balance is internally consistent."
)

SEC_5_1_BODY = [
    "Each of the four scenarios in this chapter is structured the same way. A "
    "one-sentence description of the change introduces what the desk is asked to do. "
    "Three method-narratives describe what each method does in response. A summary "
    "table compares the methods across five dimensions: what changes have to be made, "
    "how the analyst assesses the result, the rough order of analyst-hours consumed, "
    "the consistency risk level, and the kind of failure that occurs when the change "
    "is mishandled. The four scenarios in turn are: adding a grade dimension to basin "
    "production (§5.2); modelling a refinery turnaround or unplanned outage that breaks "
    "the monthly-average assumption (§5.3); receiving daily named-pipeline flow data "
    "from Genscape (§5.4); and finding inconsistencies after any of the above (§5.5). "
    "Section 5.6 aggregates the four scenarios into the headline claim.",

    "The hours and risk levels reported in the tables are rough order-of-magnitude "
    "estimates drawn from operational experience with all three method-types in "
    "commodity-analyst settings, not the result of a controlled experiment. The point "
    "of the comparison is the order-of-magnitude gap between methods on each change, "
    "not the precision of any single figure. A reader from outside oil-market analysis "
    "should be able to follow the comparisons from the tables alone; the surrounding "
    "prose provides the operational detail that makes the gap concrete.",
]

# ===================== §5.2 ADDING GRADES =====================

SEC_5_2_TITLE = "Adding grade dimensions to basin production"
SEC_5_2_BODY = [
    "The desk decides to track grades produced at each basin. WTI from the Permian, "
    "sweet light from the Eagle Ford, Bakken Light, ANS heavy, medium sour from the "
    "Gulf — each should become a tracked quantity through production, transport, "
    "storage, and refining. Aggregate barrels stay the same; the data dimension widens "
    "from {asset} to {asset, grade}.",

    ("Excel spreadsheet. ", "The existing workbook holds one column per balance line "
    "per PADD: production, runs, imports, exports, movements, stocks. Adding grades "
    "means each column duplicates by the grade dimension. Every formula must be edited. "
    "Reconciliation rules — every PADD-named-basin sum, every Jones Act in-transit "
    "allocation, every Cushing partition — must be redefined per grade. EIA per-grade "
    "reporting is sparser than aggregate, so new residual cells multiply (one per "
    "grade per PADD per balance line). The analyst's grade-allocation assumptions "
    "become embedded in formula cells, undocumented. Two analysts working from the "
    "same source data produce different grade-resolved balances because their "
    "allocation choices differ, and the differences are not visible in the totals."),

    ("Pandas notebook. ", "The DataFrame holding the balance is indexed by (date, PADD). "
    "Adding grades means widening the index to (date, PADD, grade), or pivoting to a "
    "wide format with grade-suffixed columns. Functions that transformed the aggregate "
    "balance must be rewritten to operate per grade. The new code can encapsulate "
    "allocation rules in functions rather than in cell formulas, which is more legible, "
    "but the rules themselves still have to be written and the grade definitions are "
    "bespoke to this notebook. Notebook state drift remains a risk: cells run out of "
    "order, intermediate state diverges. There is no schema-level enforcement, so a "
    "typo in a grade name silently creates a phantom grade rather than failing at the "
    "source."),

    ("Framework. ", "The asset graph stays bit-identical. The 251 nodes, the 433 "
    "directed edges, the partition tree, and the capacity ceilings do not change. "
    "What changes is one row in the commodities table (a few INSERTs) and the "
    "variables table widens from 1,870 rows to approximately 1,870 × N rows. Each "
    "new variable is either bound to a per-grade observed series (where published) "
    "or inherits the aggregate-level formula with the per-(node, grade) residual "
    "carried as B. Partition closure operates per (node, grade) automatically; "
    "v_aggregation_consistency flags any divergence per (node, grade). Mass balance "
    "closes per (node, grade) at every physical asset. The single-attribution CHECK "
    "constraint on variable_assignments applies per (variable_type, commodity, node), "
    "so a typo in a grade name fails at INSERT — it is unrepresentable."),

    "Figure 5.1 illustrates the asymmetry on a single asset — the Jones Act in-transit "
    "corridor `inter_padd_3_to_5_agg` — for visual clarity. The top panel shows each "
    "method's representation of the corridor at single-grade tracking, with the "
    "labelled stock in both columns. The bottom panel shows what happens when a grade "
    "dimension is added: the spreadsheet expands every column per grade and the "
    "in-transit row alone becomes a stack of per-grade residuals, while the framework's "
    "corridor node, edges, and capacity stay unchanged and only the variables widen.",
]

# Caption for Figure 5.1 — updated to reframe as grade-widening illustration
FIG_5_1_CAPTION_NEW = (
    "Figure 5.1. Adding a grade dimension to the balance, illustrated on a single "
    "asset. Top panel: each method's representation of the Jones Act in-transit "
    "corridor at single-grade tracking — a row in a flat table (spreadsheet) versus "
    "a node with neighbours, capacity, flow variables, and mass balance (framework). "
    "Bottom panel: what each method does when a grade dimension is added — the "
    "spreadsheet rebuilds every column per grade, while the framework keeps the same "
    "nodes and edges and widens only the variables table."
)

# Comparison Table 5.1 — Excel | Pandas | Framework × 5 dimensions
TABLE_5_1 = {
    'header': ['', 'Excel spreadsheet', 'Pandas notebook', 'Framework'],
    'rows': [
        ['What changes',
         'Cross-tab structure; every formula; reconciliation rules; source sheets',
         'DataFrame schema; every transformation; reconciliation; data loaders',
         'Add rows to commodities; data-layer pass widens variables table'],
        ['How to assess',
         'Manual sum-checks; eyeballing residuals',
         'Assertion cells; ad-hoc tests',
         'v_aggregation_consistency reports per-(node, grade) divergence'],
        ['Effort (rough)',
         '60–120 h',
         '30–60 h',
         '4–8 h'],
        ['Consistency risk',
         'High — silent disagreement between analysts',
         'Medium — visible but bespoke',
         'Low — schema-enforced'],
        ['Failure mode',
         'Bottom row reads zero; divergence buried in residuals',
         'Notebook state drift; coverage gaps in assertions',
         'Inconsistency unrepresentable for constrained patterns'],
    ],
    'caption': (
        "Table 5.1. Adding a grade dimension to basin production: comparison across "
        "three methods. The framework's asset graph stays unchanged; only the variables "
        "table widens by the number of grades. The spreadsheet and the notebook rebuild "
        "to varying degrees and run a higher consistency risk."
    ),
}

# ===================== §5.3 TAR / OUTAGE =====================

SEC_5_3_TITLE = "Refinery turnaround or unplanned outage"
SEC_5_3_BODY = [
    "A refinery enters an unplanned outage on day 5 of February and resumes operations "
    "on day 28. For twenty-three days the refinery's crude runs are zero; for the other "
    "five days they are at the design rate. The monthly EIA aggregate, which arrives at "
    "the end of March, reports the month-average runs — about 15% of nominal capacity. "
    "The desk needs to track the actual step-function in real time, because the question "
    "is not what February's average was; it is which Gulf Coast tank farms are filling up "
    "because their downstream offtake stopped, and whether crude vessels destined for the "
    "refinery were diverted to alternative receivers.",

    ("Excel spreadsheet. ", "The existing workbook is structured around monthly EIA "
    "series. The cells assume monthly resolution: one row per month, one column per asset. "
    "To track the daily step-function, the analyst must either build a parallel daily "
    "sheet for the affected refinery and its upstream tank farms — manually reconciling "
    "the daily figures with the EIA monthly aggregate when it arrives at month-end — or "
    "abandon the daily detail and accept that the monthly figure represents an average "
    "that misrepresents what actually happened. The first option is operationally correct; "
    "the second is what most workbooks default to. If the analyst chooses the first, the "
    "propagation to upstream tank farms is also manual: an inventory-accumulation formula "
    "on each affected tank farm sheet, with the rate change happening at the day of the "
    "outage and reversing at the day of the restart. Capacity ceilings are not checked "
    "automatically."),

    ("Pandas notebook. ", "The notebook holds the balance as a DataFrame indexed by date. "
    "Switching one asset to daily resolution while others remain monthly is technically "
    "straightforward — resample('M') aggregates the daily back to monthly, and joins "
    "between monthly and daily series can be done with pandas' alignment semantics. But "
    "the analyst still has to decide the resampling rule (mean? sum?), handle missing days "
    "(forward-fill? interpolation? raise an error?), and manage mixed-resolution joins "
    "where a monthly series is being compared against a daily-aggregated sum. The "
    "propagation logic must be written: if refinery X is offline for twenty-three days, "
    "what does that imply for tank farm Y's inventory? Mass balance is an assertion in "
    "code, not a structural property of the data model."),

    ("Framework. ", "The timeseries_data table already carries observations at any "
    "granularity — daily, weekly, monthly observations coexist in the same table. For the "
    "refinery TAR, the analyst binds a daily IIR Energy series to the refinery's "
    "consumption variable for the duration of the outage; the resolver dispatches the "
    "daily observations directly. Aggregation to the monthly figure is handled by the "
    "resolver's mean-over-window rule. Mass balance at the refinery node holds at the "
    "daily resolution, so the resolver propagates the consumption drop to the upstream "
    "tank farm's inventory accumulation automatically through the partition closure on "
    "the tank farm's mass-balance equation. The capacity ceiling on the tank farm is "
    "asset_capacities.capacity_bbl, and v_capacity_violations reports any day on which "
    "the implied tank farm inventory exceeds capacity. The single-attribution CHECK on "
    "variable_assignments ensures the daily series and the monthly EIA aggregate are "
    "not both bound to the same variable in the same scenario."),

    "Hurricane Harvey, August 2017, provides a documented example. The storm made "
    "landfall on the Texas Gulf Coast on 25 August 2017 as a Category 4 hurricane, and "
    "approximately 25% of U.S. refining capacity was offline at peak for periods ranging "
    "from several days to over two weeks. The PADD 3 monthly EIA balance for August 2017 "
    "produced an implied residual of +193 kbd that the spreadsheet workflow absorbs into "
    "an unlabelled adjustment cell and the framework retains as a labelled balancing item "
    "B on the PADD 3 node. The residual is consistent with crude that left PADD 3 by "
    "tanker before the storm and was reported as exported, while the receiving inventories "
    "at downstream destinations had not yet caught up; with delayed pipeline "
    "custody-transfer reports; and with refinery process losses that fell as throughput "
    "collapsed. Section 5.5 returns to this point with a table of comparable spikes around "
    "COVID-19 (March 2020), the Colonial Pipeline shutdown (May 2021), and the Strategic "
    "Petroleum Reserve release (May 2022 peak).",
]

TABLE_5_2 = {
    'header': ['', 'Excel spreadsheet', 'Pandas notebook', 'Framework'],
    'rows': [
        ['What changes',
         'New daily sheets; reconciliation rules; inventory propagation formulas per asset',
         'DataFrame schema; resampling logic; propagation functions',
         'One INSERT into variable_assignments binding a daily series'],
        ['How to assess',
         'Manual sum-checks; eyeball inventory trajectories',
         'Assertion suite; resampling-consistency checks',
         'v_node_balance_check; v_capacity_violations'],
        ['Effort (rough)',
         '40–80 h setup + 4–8 h per event',
         '20–40 h setup + 2–4 h per event',
         '2–4 h per event'],
        ['Consistency risk',
         'High — daily/monthly drift; analyst-specific reconciliation rules',
         'Medium — visible but bespoke; mixed-resolution semantics drift',
         'Low — mass balance enforced per node per day; capacity ceilings queryable'],
        ['Failure mode',
         'Inventory ceiling breaches go undetected; reconciliation rules diverge between analysts',
         'False-negative assertions; resampling-rule choices undocumented',
         'Inconsistency unrepresentable for constrained patterns'],
    ],
    'caption': (
        "Table 5.2. Modelling a refinery turnaround or unplanned outage: comparison "
        "across three methods. The framework binds the daily series to an existing "
        "node; propagation to upstream assets is automatic via mass balance. The "
        "spreadsheet and the notebook require setup work per asset plus per-event "
        "effort, with consistency enforcement left to the analyst."
    ),
}

# ===================== §5.4 GENSCAPE =====================

SEC_5_4_TITLE = "Receiving Genscape pipeline data"
SEC_5_4_BODY = [
    "The desk subscribes to Wood Mackenzie's Genscape service, which reports daily flow "
    "rates on named pipelines via electromagnetic sensors mounted on the lines themselves. "
    "Approximately 91% of capacity connected to the Cushing hub is covered, including "
    "the major Permian outbound (Gray Oak, Cactus II, BridgeTex, Midland-to-ECHO), the "
    "Cushing connectors (Seaway, Spearhead, Pony Express), and the U.S.-Canada border "
    "pipelines (Enbridge Mainline, Keystone). EIA continues to report only aggregate "
    "inter-PADD movements; Genscape gives a daily breakdown for the named pipelines.",

    ("Excel spreadsheet. ", "Adding Genscape data to a working EIA spreadsheet looks "
    "straightforward on first inspection — append columns, point the formulas at the new "
    "cells, recompute the balance — but the integration breaks on five distinct mismatches "
    "that the spreadsheet has no schema-level way to resolve. The frequencies disagree: "
    "Genscape reports daily, EIA monthly, and a pipeline that is down for four days "
    "mid-month leaves no trace in its monthly mean. The geographies disagree: Genscape "
    "names a pipeline (DAPL, Keystone), while EIA names a PADD pair, and a single EIA "
    "inter-PADD figure aggregates over an undisclosed set of pipelines, some of which "
    "Genscape names individually and some of which it does not. Double-counting becomes "
    "a live risk on every new column: if DAPL flows are appended as a separate series and "
    "the EIA PADD-2-to-PADD-3 movement series already incorporates them, every barrel on "
    "DAPL is now in the balance twice unless the analyst manually subtracts it from one "
    "side. Coverage is incomplete: the nine percent of Cushing-connected capacity that "
    "Genscape does not monitor still shows up in the EIA aggregate, so a "
    "residual-pipeline column has to be maintained by hand each month. And the "
    "integration is not stable: every time Genscape adds coverage of a new pipeline "
    "(Seminole-Red, EPIC Crude, the latest reversed segment), the analyst must add a "
    "column, decide which existing aggregate it should be subtracted from, and update the "
    "residual definition. None of these decisions are encoded; they live in the analyst's "
    "head and in the workbook's commented cells, and they evolve as personnel rotate."),

    ("Pandas notebook. ", "The notebook can ingest Genscape's API as a DataFrame and join "
    "with EIA aggregates. The structural advantages over Excel are real: the join logic "
    "is explicit code rather than hidden cell formulas; assertions can check that "
    "subtractions are applied; a function can encode the 'Genscape covers pipelines A, B, "
    "C; EIA aggregate is over A, B, C, D, E' mapping. But each new pipeline addition still "
    "requires updating the mapping function; the double-counting check is an assertion, "
    "not a structural property; mixed-resolution semantics (daily vs monthly) must be "
    "coded explicitly. As Genscape's coverage grows, the mapping function and the "
    "assertion suite must grow with it."),

    ("Framework. ", "Genscape data binds to existing pipeline nodes. For each pipeline "
    "Genscape names (DAPL, Keystone, Gray Oak, Cactus II, …), the framework already has "
    "a node in the asset graph; the daily flow series binds to that node's outflow "
    "variable. The aggregation view ties Genscape-bound pipeline flows to the PADD-view "
    "inflow/outflow variables through the partition formula on the PADD-view, so the EIA "
    "inter-PADD aggregate is automatically reconciled against the sum of named pipeline "
    "flows. The 9% coverage gap emerges as the difference on the PADD-view node "
    "automatically, carried as B at the PADD-view. The schema-level CHECK on single "
    "attribution prevents the double-counting failure mode: the same variable cannot be "
    "bound to both the Genscape series and the EIA aggregate; one is the source of value, "
    "the other is the constraint set against which v_aggregation_consistency reports "
    "divergence. For new pipeline coverage additions: add a TS binding on the existing "
    "pipeline node; the aggregation view picks it up automatically; the residual on the "
    "PADD-view rebalances automatically; no other code changes are required."),
]

TABLE_5_3 = {
    'header': ['', 'Excel spreadsheet', 'Pandas notebook', 'Framework'],
    'rows': [
        ['What changes',
         'Columns per pipeline; manual subtractions; residual definition',
         'Join code; pipeline-to-aggregate mapping function; assertions',
         'N INSERTs into variable_assignments (one per Genscape-named pipeline)'],
        ['How to assess',
         'Manual sum-checks; eyeball Genscape totals against EIA',
         'Assertion suite; plots of Genscape sum vs EIA aggregate',
         'v_aggregation_consistency reports divergence per (PADD, month)'],
        ['Effort initial',
         '20–40 h',
         '15–30 h',
         'N × 1–2 h'],
        ['Effort per new pipeline',
         '2–4 h',
         '1–3 h',
         '1–2 h'],
        ['Consistency risk',
         'High — silent double-counting; residual-definition drift',
         'Medium — bespoke assertions; mapping-function staleness',
         'Low — schema-enforced single attribution; coverage gap auto-flagged'],
    ],
    'caption': (
        "Table 5.3. Integrating Genscape daily pipeline-flow data on top of EIA monthly "
        "aggregates: comparison across three methods. The framework binds each Genscape "
        "series to an existing pipeline node; the aggregation view reconciles "
        "automatically, and double-counting is unrepresentable. The spreadsheet and the "
        "notebook accumulate per-pipeline integration debt."
    ),
}

# ===================== §5.5 FINDING INCONSISTENCIES =====================

SEC_5_5_TITLE = "Finding inconsistencies"
SEC_5_5_BODY = [
    "After a month-end refresh of the balance — whether from monthly EIA, daily Genscape, "
    "or any combination — the desk needs to verify that the balance is internally "
    "consistent. Did all PADD-named-basin sums tie to PADD totals? Did all pipeline flows "
    "balance against inter-PADD aggregates? Did inventory changes match flow differences? "
    "Did any flow exceed its asset capacity? Did any variable get double-bound to two "
    "conflicting series? These are not exotic questions; they are the daily verification "
    "a desk does before relying on the balance.",

    ("Excel spreadsheet. ", "Inconsistency-finding in a spreadsheet is by manual "
    "inspection. The analyst writes one or more check sheets where critical sums are "
    "recomputed and the difference is displayed in a corner cell, often colour-coded if "
    "it exceeds a tolerance. The coverage is whatever the analyst remembers to check. A "
    "typical sequence is: PADD totals against named-basin sums; monthly stock changes "
    "against (production − consumption + imports − exports + movements_in − "
    "movements_out); refinery runs against the EIA refining-district aggregate. Outside "
    "the analyst's check sheets, inconsistencies are silent. When the workbook is handed "
    "over to a new analyst, the check coverage walks away with the previous one."),

    ("Pandas notebook. ", "Inconsistency-finding in a notebook is more systematic. The "
    "analyst writes assertion cells: assert (padd_sum == padd_total).all(); assert "
    "(stock_change == flows.sum()).all(). These assertions can be packaged as a test "
    "suite and run automatically on each balance refresh. Coverage is whatever the "
    "analyst writes — better than a spreadsheet's manual sheets but still bespoke. The "
    "Achilles' heel of assertion-based testing is false negatives: an assertion that's "
    "never run never catches anything. As the scenarios evolve — new grades, new "
    "pipelines, new outage models — the assertion suite must be maintained to keep pace. "
    "Notebook state drift means the same cell can produce different assertion outcomes "
    "across re-runs."),

    ("Framework. ", "Inconsistency-finding in the framework runs in two layers: "
    "schema-level enforcement at write time, and view-level audit at read time. The "
    "schema layer carries the CHECK constraint on num_nonnulls(timeseries_id, formula) "
    "= 1 on variable_assignments, foreign-key constraints across the asset graph, and "
    "trigger-enforced same-graph rules. The view layer materialises three consistency "
    "audits that re-run automatically when the resolver completes:"),

    "v_aggregation_consistency reports any (node, variable_type) pair where the observed "
    "parent value disagrees with the sum of partition children, per observation date, with "
    "the magnitude of the disagreement quantified. v_resolution_anomalies reports "
    "resolver-level oddities — long forward-fill runs, negative derived values, partial "
    "dispatch (some children observed, some latent) — with severity ranked. "
    "v_capacity_violations reports any flow that exceeds the receiving or originating "
    "asset's capacity_bd ceiling.",

    "The schema enforcement is part of the database; the audit views are part of the "
    "resolver pipeline. Neither requires per-scenario setup. New grades, new pipelines, "
    "new outage observations — all are checked automatically by the same views without "
    "any change to the audit infrastructure. The framework's B series at the USA "
    "aggregate level also acts as a passive monitor of where the reporting system is and "
    "is not closing cleanly. Table 5.5 reports the absolute value of B at the USA "
    "aggregate around four named disruption events, alongside the ten-year median for "
    "comparison.",
]

# Comparison table 5.4 + named-events table 5.5
TABLE_5_4 = {
    'header': ['', 'Excel spreadsheet', 'Pandas notebook', 'Framework'],
    'rows': [
        ['What changes per scenario',
         'New check sheets; tolerance tuning',
         'New assertion cells; suite maintenance',
         'Nothing — built into schema + resolver pipeline'],
        ['How to assess',
         'Read check sheets visually',
         'Run test suite; inspect failures',
         'Query v_aggregation_consistency, v_resolution_anomalies, v_capacity_violations'],
        ['Effort initial',
         '5–20 h',
         '10–30 h',
         '0 h (built-in)'],
        ['Effort per scenario change',
         '1–2 h per month',
         '2–5 h per scenario',
         '0 h'],
        ['Consistency risk',
         'High — silent coverage gaps; analyst-specific tolerances',
         'Medium — false-negative assertions; coverage gaps where assertions are missing',
         'Low — schema-level CHECKs + automatic audit views'],
    ],
    'caption': (
        "Table 5.4. Finding inconsistencies after a balance refresh: comparison across "
        "three methods. The framework's consistency-audit infrastructure is built into "
        "the schema and the resolver pipeline; it requires zero per-scenario setup. The "
        "spreadsheet and the notebook require per-scenario test infrastructure that the "
        "analyst maintains."
    ),
}

TABLE_5_5_NAMED_EVENTS = {
    'header': ['Event', 'Month', '|B| at USA agg (kbd)', 'Multiple of median'],
    'rows': [
        ['Median, all months 2015–2024 (baseline)', '—', '~260', '1.0×'],
        ['Hurricane Harvey', 'Aug 2017', '~145', '0.6×'],
        ['COVID-19 demand collapse', 'Mar 2020', '~748', '2.9×'],
        ['Colonial Pipeline shutdown', 'May 2021', '~596', '2.3×'],
        ['Strategic Petroleum Reserve release', 'May 2022 (peak)', '~427', '1.7×'],
    ],
    'caption': (
        "Table 5.5. Absolute value of the balancing item B at the USA aggregate around "
        "named disruption events, compared with the ten-year median. The framework "
        "retains B as a first-class labelled variable, so spikes around named events are "
        "visible in the resolved output without bespoke audit code. The spreadsheet "
        "workflow absorbs the same residual into an unlabelled adjustment cell. "
        "Hurricane Harvey at the USA aggregate is the partial exception: its national "
        "signal is partly cancelled by offsetting flows between PADD 3 and the rest of "
        "the country, while the per-PADD signal documented in §5.3 remains clear."
    ),
}

# ===================== §5.6 SYNTHESIS =====================

SEC_5_6_TITLE = "What the comparison establishes"
SEC_5_6_BODY = [
    "The four scenarios establish a consistent pattern. The cost of evolving the balance, "
    "measured in analyst-hours, differs by one or two orders of magnitude between the "
    "framework and either of the alternatives, summarised in Table 5.6.",

    "What drives the gap is not technical sophistication. A competent Python notebook "
    "can do everything a spreadsheet can do and more — modular code, version control, "
    "programmatic assertions, mixed-resolution joins. The notebook's advantage over the "
    "spreadsheet on every scenario is genuine. The framework's advantage over the "
    "notebook is genuine in turn, but it has a different character: it is structural "
    "rather than operational. The framework enforces single attribution, partition "
    "closure, and capacity constraints at the database level, before any code runs. The "
    "notebook can replicate these checks as assertions, but the assertions are bespoke "
    "per scenario, depend on the analyst remembering to write them, and can be silently "
    "bypassed when the notebook is re-run with stale state.",

    "The cost asymmetry compounds because the four scenarios are rarely encountered in "
    "isolation. A real desk does all four changes in a single quarter — adds grade "
    "tracking for a specific basin, models a Gulf Coast TAR, integrates Genscape "
    "coverage on the Permian outbound complex, and runs consistency audits monthly. The "
    "framework's setup cost is one-time and structural; the alternatives' setup cost is "
    "per-scenario and additive. Over a year of operational use, the gap is not 60–120 "
    "hours versus 4–8 hours on grades; it is the sum across all scenarios that the desk "
    "encounters, with consistency risk compounding silently in the alternatives.",

    "The framework is not a methodologically more sophisticated balance — it produces "
    "the same monthly aggregate values from the same EIA series in steady state. It is "
    "a representation that absorbs change without requiring rebuild, and that enforces "
    "consistency at the schema level rather than relying on analyst vigilance. The "
    "thesis claim is that this is the structural property a representation should have "
    "for production use in a setting where balances evolve continuously and consistency "
    "failures are operationally costly.",
]

TABLE_5_6_AGGREGATE = {
    'header': ['Scenario', 'Excel spreadsheet', 'Pandas notebook', 'Framework'],
    'rows': [
        ['1. Adding grade dimensions to basin production',
         '60–120 h',
         '30–60 h',
         '4–8 h'],
        ['2. Refinery TAR / unplanned outage',
         '40–80 h + 4–8 h per event',
         '20–40 h + 2–4 h per event',
         '2–4 h per event'],
        ['3. Receiving Genscape pipeline data',
         '20–40 h initial; 2–4 h per pipeline added',
         '15–30 h initial; 1–3 h per pipeline added',
         '1–2 h per pipeline added'],
        ['4. Finding inconsistencies',
         '5–20 h initial; 1–2 h per month',
         '10–30 h initial; 2–5 h per scenario',
         '0 h (built-in)'],
    ],
    'caption': (
        "Table 5.6. Aggregate cost asymmetry across the four change-management scenarios. "
        "Hours are rough order-of-magnitude estimates drawn from operational experience. "
        "The point of the table is the order-of-magnitude gap between methods, not the "
        "precision of any single figure. The framework's setup cost is one-time and "
        "structural; the alternatives' setup cost is per-scenario and additive over the "
        "operational lifetime of the balance."
    ),
}


# =================== HELPERS ===================

def make_para(text, *, italic=False, bold=False, bold_lead=None):
    """Create a w:p element with text. If bold_lead is given, that prefix is bolded."""
    p = OxmlElement('w:p')
    if bold_lead:
        r1 = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rPr.append(OxmlElement('w:b'))
        r1.append(rPr)
        t1 = OxmlElement('w:t')
        t1.text = bold_lead
        t1.set(qn('xml:space'), 'preserve')
        r1.append(t1)
        p.append(r1)
        r2 = OxmlElement('w:r')
        t2 = OxmlElement('w:t')
        t2.text = text
        t2.set(qn('xml:space'), 'preserve')
        r2.append(t2)
        p.append(r2)
    else:
        r = OxmlElement('w:r')
        if italic or bold:
            rPr = OxmlElement('w:rPr')
            if bold:
                rPr.append(OxmlElement('w:b'))
            if italic:
                rPr.append(OxmlElement('w:i'))
            r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
    return p


def _add_table_borders(table_el):
    """Add single-line borders to all sides + insideH/V of the table."""
    tblPr = table_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table_el.insert(0, tblPr)
    # Remove any existing borders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '666666')
        borders.append(b)
    tblPr.append(borders)


def make_table(doc, data, width_inches=6.5):
    """Create a w:tbl with the given data. data: {'header': [str], 'rows': [[str]], 'caption': str}."""
    n_cols = len(data['header'])
    n_rows = 1 + len(data['rows'])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    # No table style — apply borders manually so it renders visibly
    _add_table_borders(table._element)
    # Header row
    for j, h in enumerate(data['header']):
        cell = table.cell(0, j)
        cell.text = h
        # bold the header text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    # Body rows
    for i, row in enumerate(data['rows'], start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = val
    return table


def find_and_replace_text(element, old, new):
    n = 0
    for t in element.findall('.//' + qn('w:t')):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            n += 1
    return n


def set_para_text(para_el, text):
    """Replace para's runs with a single run containing `text`."""
    for r in para_el.findall(qn('w:r')):
        para_el.remove(r)
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    para_el.append(new_r)


def find_heading(body, heading_text, style='Heading2'):
    """Find the first Heading2 paragraph whose text contains heading_text (substring match)."""
    for el in body:
        if el.tag != qn('w:p'):
            continue
        pStyle = el.find('.//' + qn('w:pStyle'))
        if pStyle is None:
            continue
        if pStyle.get(qn('w:val')) != style:
            continue
        texts = el.findall('.//' + qn('w:t'))
        text = ''.join(t.text or '' for t in texts).strip()
        if heading_text in text:
            return el
    return None


def collect_section_body(body, heading_el, next_heading_styles=('Heading2','Heading1')):
    """Collect all body elements after heading_el up to (but not including) the next heading."""
    out = []
    found = False
    for el in body:
        if el is heading_el:
            found = True
            continue
        if not found:
            continue
        if el.tag == qn('w:p'):
            pStyle = el.find('.//' + qn('w:pStyle'))
            if pStyle is not None and pStyle.get(qn('w:val')) in next_heading_styles:
                break
        out.append(el)
    return out


# =================== MAIN ===================

def main():
    doc = Document(str(SRC))
    body = doc.element.body

    # === Step 1: locate section headings ===
    sections = {}
    for label, hdr_substring in [
        ('5.1', '5.1  Setup'),
        ('5.2', 'Jones Act in-transit anomaly'),  # current §5.2
        ('5.3', 'Stress-test comparison: Hurricane Harvey'),  # current §5.3
        ('5.4', 'Multi-source integration'),  # current §5.4
        ('5.5', 'The same pattern on other named events'),  # current §5.5
        ('5.6', 'What the comparison establishes'),  # current §5.6
    ]:
        h = find_heading(body, hdr_substring)
        if h is None:
            print(f"FAIL: section {label} heading not found ({hdr_substring!r})")
            return 1
        sections[label] = h
        text = ''.join(t.text or '' for t in h.findall('.//' + qn('w:t'))).strip()
        print(f"  found §{label}: {text[:90]}")

    # Also find chapter 5 heading + chapter intro paragraph + next chapter heading (Conclusion)
    chapter5_hdr = find_heading(body, 'Case study: where the framework is better', style='Heading1')
    chapter6_hdr = find_heading(body, 'Conclusion and future work', style='Heading1')
    if chapter5_hdr is None or chapter6_hdr is None:
        print("FAIL: chapter boundaries not located")
        return 1
    print("  found chapter boundaries")

    # === Step 2: delete bodies of all sections (keep heading + Figure 5.1 paragraph for §5.2) ===
    # Locate Figure 5.1 paragraph (one with <w:drawing>)
    figure_para = None
    for el in body:
        if el.tag != qn('w:p'):
            continue
        if el.findall('.//' + qn('w:drawing')):
            # this might be a figure; check if it's the Figure 5.1 in §5.2
            figure_para = el  # we'll detect later by position
            # for now just take the first figure in chapter 5 region
            # but we want the figure inside §5.2 — let's collect all figures and pick the §5.2 one

    # Better: collect all figures and check which is in chapter 5
    figure_paras_in_ch5 = []
    in_ch5 = False
    for el in body:
        if el is chapter5_hdr:
            in_ch5 = True
            continue
        if el is chapter6_hdr:
            in_ch5 = False
            break
        if in_ch5 and el.tag == qn('w:p') and el.findall('.//' + qn('w:drawing')):
            figure_paras_in_ch5.append(el)

    if len(figure_paras_in_ch5) != 1:
        print(f"WARN: found {len(figure_paras_in_ch5)} figure paragraphs in chapter 5")
    if figure_paras_in_ch5:
        figure_para = figure_paras_in_ch5[0]
        print(f"  identified Figure 5.1 paragraph")

    # === Step 2a: rewrite the chapter intro paragraph (the one right after chapter 5 heading) ===
    # Find paragraph immediately following chapter5_hdr
    seen = False
    chapter5_intro_para = None
    for el in body:
        if el is chapter5_hdr:
            seen = True
            continue
        if seen and el.tag == qn('w:p'):
            chapter5_intro_para = el
            break
    if chapter5_intro_para is not None:
        set_para_text(chapter5_intro_para, CHAPTER5_INTRO)
        print(f"  rewrote chapter intro paragraph")
    else:
        print("WARN: could not find chapter 5 intro paragraph")

    # === Step 2b: collect body elements of each section (everything between heading and next heading) ===
    # Order: §5.1 → §5.2 → §5.3 → §5.4 → §5.5 → §5.6 → chapter6
    section_order = ['5.1','5.2','5.3','5.4','5.5','5.6']
    next_boundaries = {
        '5.1': sections['5.2'],
        '5.2': sections['5.3'],
        '5.3': sections['5.4'],
        '5.4': sections['5.5'],
        '5.5': sections['5.6'],
        '5.6': chapter6_hdr,
    }

    # For each section, delete all body elements between heading and boundary
    # EXCEPT for Figure 5.1 in §5.2 (we'll keep it and just update the surrounding content + caption)

    body_list = list(body)
    for sec in section_order:
        hdr = sections[sec]
        boundary = next_boundaries[sec]
        # collect elements to delete
        to_delete = []
        in_section = False
        for el in body_list:
            if el is hdr:
                in_section = True
                continue
            if el is boundary:
                break
            if not in_section:
                continue
            # don't delete the figure in §5.2 — we keep it
            if sec == '5.2' and el is figure_para:
                continue
            to_delete.append(el)
        for el in to_delete:
            body.remove(el)
        print(f"  §{sec}: deleted {len(to_delete)} body elements")

    # === Step 3: update each section's heading text and insert new content ===

    # --- §5.1 Setup ---
    # Heading text stays (it's "5.1  Setup"); insert body paragraphs
    prev = sections['5.1']
    for text in SEC_5_1_BODY:
        new_p = make_para(text)
        prev.addnext(new_p)
        prev = new_p
    print(f"  §5.1: inserted {len(SEC_5_1_BODY)} body paragraphs")

    # --- §5.2 Adding grades ---
    find_and_replace_text(sections['5.2'],
                          'Steady-state comparison: the Jones Act in-transit anomaly',
                          SEC_5_2_TITLE)
    prev = sections['5.2']
    for item in SEC_5_2_BODY:
        if isinstance(item, tuple):
            bold_lead, body_text = item
            new_p = make_para(body_text, bold_lead=bold_lead)
        else:
            new_p = make_para(item)
        prev.addnext(new_p)
        prev = new_p
    # Now figure 5.1 should already be in the body but it's been "kept" before prev.
    # We need to position the figure at the end of the §5.2 body (or wherever sensible).
    # Actually — at this point, the figure_para is still in its original position in the body,
    # which is BEFORE all the new paragraphs we just inserted (because the figure was preserved
    # while everything around it was deleted, and then we inserted new content AFTER the heading).
    # We need to MOVE the figure to after the last §5.2 body paragraph + caption.
    if figure_para is not None:
        # Detach figure from its current position
        figure_para.getparent().remove(figure_para)
        # Re-insert at the end of §5.2 body (after the closing paragraph just inserted as `prev`)
        prev.addnext(figure_para)
        prev = figure_para
    # Update figure caption
    # Find caption paragraph — should be the paragraph immediately after figure (if it survived)
    # But we deleted all non-figure paragraphs, so we need to add a new caption
    new_caption = make_para(FIG_5_1_CAPTION_NEW, italic=True)
    prev.addnext(new_caption)
    prev = new_caption
    # Insert Table 5.1
    new_table_5_1 = make_table(doc, TABLE_5_1)
    # Move table to position
    table_el = new_table_5_1._element
    table_el.getparent().remove(table_el)
    prev.addnext(table_el)
    prev = table_el
    # Caption after table
    caption_5_1 = make_para(TABLE_5_1['caption'], italic=True)
    prev.addnext(caption_5_1)
    print(f"  §5.2: rewrote title, inserted body, repositioned figure, added Table 5.1")

    # --- §5.3 TAR/outage ---
    find_and_replace_text(sections['5.3'],
                          'Stress-test comparison: Hurricane Harvey, August 2017',
                          SEC_5_3_TITLE)
    prev = sections['5.3']
    for item in SEC_5_3_BODY:
        if isinstance(item, tuple):
            new_p = make_para(item[1], bold_lead=item[0])
        else:
            new_p = make_para(item)
        prev.addnext(new_p)
        prev = new_p
    # Table 5.2
    new_table_5_2 = make_table(doc, TABLE_5_2)
    t = new_table_5_2._element
    t.getparent().remove(t)
    prev.addnext(t)
    prev = t
    prev.addnext(make_para(TABLE_5_2['caption'], italic=True))
    print(f"  §5.3: rewrote title, inserted body, added Table 5.2")

    # --- §5.4 Genscape ---
    find_and_replace_text(sections['5.4'],
                          'Multi-source integration: layering pipeline-flow data on top of EIA aggregates',
                          SEC_5_4_TITLE)
    prev = sections['5.4']
    for item in SEC_5_4_BODY:
        if isinstance(item, tuple):
            new_p = make_para(item[1], bold_lead=item[0])
        else:
            new_p = make_para(item)
        prev.addnext(new_p)
        prev = new_p
    # Table 5.3
    new_table_5_3 = make_table(doc, TABLE_5_3)
    t = new_table_5_3._element
    t.getparent().remove(t)
    prev.addnext(t)
    prev = t
    prev.addnext(make_para(TABLE_5_3['caption'], italic=True))
    print(f"  §5.4: rewrote title, inserted body, added Table 5.3")

    # --- §5.5 Finding inconsistencies ---
    find_and_replace_text(sections['5.5'],
                          'The same pattern on other named events',
                          SEC_5_5_TITLE)
    prev = sections['5.5']
    for item in SEC_5_5_BODY:
        if isinstance(item, tuple):
            new_p = make_para(item[1], bold_lead=item[0])
        else:
            new_p = make_para(item)
        prev.addnext(new_p)
        prev = new_p
    # Table 5.4 (comparison)
    new_table_5_4 = make_table(doc, TABLE_5_4)
    t = new_table_5_4._element
    t.getparent().remove(t)
    prev.addnext(t)
    prev = t
    prev.addnext(make_para(TABLE_5_4['caption'], italic=True))
    prev = list(body)[-1]  # the caption we just added
    # Actually we need to advance prev properly — let's find it
    # The caption is the last paragraph appended; its element is at the end now
    # But it was inserted via addnext(), so it's right after the table
    # Re-resolve prev to be the caption we just added
    # Walk from §5.5 heading forward to find the last paragraph in §5.5 region
    last_in_5_5 = None
    in_5_5 = False
    for el in body:
        if el is sections['5.5']:
            in_5_5 = True
            continue
        if el is sections['5.6']:
            break
        if in_5_5:
            last_in_5_5 = el
    prev = last_in_5_5
    # Table 5.5 (named events)
    new_table_5_5 = make_table(doc, TABLE_5_5_NAMED_EVENTS)
    t = new_table_5_5._element
    t.getparent().remove(t)
    prev.addnext(t)
    prev = t
    prev.addnext(make_para(TABLE_5_5_NAMED_EVENTS['caption'], italic=True))
    print(f"  §5.5: rewrote title, inserted body, added Tables 5.4 + 5.5")

    # --- §5.6 Synthesis ---
    # Title stays
    prev = sections['5.6']
    for text in SEC_5_6_BODY:
        new_p = make_para(text)
        prev.addnext(new_p)
        prev = new_p
    # Table 5.6 aggregate cost
    new_table_5_6 = make_table(doc, TABLE_5_6_AGGREGATE)
    t = new_table_5_6._element
    t.getparent().remove(t)
    prev.addnext(t)
    prev = t
    prev.addnext(make_para(TABLE_5_6_AGGREGATE['caption'], italic=True))
    print(f"  §5.6: inserted body + aggregate cost Table 5.6")

    # === Step 4: update TOC entries (front-matter static text) ===
    # TOC entries are like "5.2   Steady-state comparison: the Jones Act in-transit anomaly23"
    toc_renames = [
        ('Steady-state comparison: the Jones Act in-transit anomaly', SEC_5_2_TITLE),
        ('Stress-test comparison: Hurricane Harvey, August 2017', SEC_5_3_TITLE),
        ('Multi-source integration: layering pipeline-flow data on top of EIA aggregates', SEC_5_4_TITLE),
        ('The same pattern on other named events', SEC_5_5_TITLE),
    ]
    # Apply by direct XML find-replace on document.xml — easier than walking docx tree for TOC
    # python-docx doesn't expose TOC entries cleanly; do via lxml on body
    for old_title, new_title in toc_renames:
        for el in body:
            if el.tag != qn('w:p'):
                continue
            for t in el.findall('.//' + qn('w:t')):
                if t.text and old_title in t.text:
                    t.text = t.text.replace(old_title, new_title)

    # === Step 5: update List of Tables entries ===
    # LoT entries reference the table captions; update them to match new captions
    # The old LoT was: Table 5.1 (PADD 5 stock identity), 5.2 (PADD 3 monthly balance Aug 2017),
    #                  5.3 (Genscape layering), 5.4 (Named events)
    # New LoT should be: 5.1 (Adding grades), 5.2 (Refinery TAR), 5.3 (Genscape), 5.4 (Finding inconsistencies),
    #                    5.5 (Named events), 5.6 (Aggregate cost)
    lot_renames = [
        # (old_substring, new_text)
        ('PADD 5 stock identity, representative month',
         'Adding a grade dimension to basin production: comparison across three methods'),
        ('PADD 3 monthly balance, August 2017',
         'Modelling a refinery turnaround or unplanned outage: comparison across three methods'),
        ('Layering Genscape pipeline-flow data on top of EIA monthly aggregates',
         'Integrating Genscape daily pipeline-flow data: comparison across three methods'),
        ('Absolute value of B at the USA aggregate around named disruption events',
         'Finding inconsistencies after a balance refresh: comparison across three methods'),
    ]
    for old, new in lot_renames:
        for el in body:
            if el.tag != qn('w:p'):
                continue
            for t in el.findall('.//' + qn('w:t')):
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new)

    # Add new LoT entries for Tables 5.5 and 5.6
    # Find the last LoT entry (currently Table 5.4) and clone it for 5.5 and 5.6
    last_lot = None
    for el in body:
        if el.tag != qn('w:p'):
            continue
        texts = el.findall('.//' + qn('w:t'))
        text = ''.join(t.text or '' for t in texts).strip()
        if text.startswith('Table 5'):
            last_lot = el
    if last_lot is not None:
        # Clone for Table 5.5
        new_lot_5_5 = deepcopy(last_lot)
        ts = new_lot_5_5.findall('.//' + qn('w:t'))
        if ts:
            ts[0].text = ('Table 5.5   Absolute value of B at the USA aggregate '
                          'around named disruption events')
        for i in range(1, len(ts)-1):
            ts[i].text = ''
        if len(ts) >= 1:
            ts[-1].text = '32'  # placeholder page
        last_lot.addnext(new_lot_5_5)

        # Clone for Table 5.6
        new_lot_5_6 = deepcopy(last_lot)
        ts = new_lot_5_6.findall('.//' + qn('w:t'))
        if ts:
            ts[0].text = ('Table 5.6   Aggregate cost asymmetry across the four '
                          'change-management scenarios')
        for i in range(1, len(ts)-1):
            ts[i].text = ''
        if len(ts) >= 1:
            ts[-1].text = '33'  # placeholder page
        new_lot_5_5.addnext(new_lot_5_6)

    # === Save ===
    doc.save(str(DST))
    print(f"\nSaved: {DST.name}")
    print(f"  source size: {SRC.stat().st_size:,} bytes")
    print(f"  output size: {DST.stat().st_size:,} bytes")
    return 0


if __name__ == '__main__':
    import sys
    sys.exit(main())
