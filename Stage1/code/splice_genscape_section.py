"""One-shot: splice a new §5.4 "Multi-source integration" worked example into
Master_Thesis_Pedro_Porfirio_restructured_v2.2.docx, renumber existing §5.4
and §5.5, update Table 5.3 cross-references, write v2.3.

The new section is the Genscape vs IEA spreadsheet case study — a third case
study in Chapter 5 alongside PADD 5 in-transit (§5.2) and Hurricane Harvey
(§5.3). Adds Table 5.3 (Genscape comparison) and pushes the previous Table 5.3
to Table 5.4.
"""
from __future__ import annotations
from pathlib import Path
from copy import deepcopy

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL

SRC = Path(__file__).resolve().parent.parent / "outputs" / "docs" / "Master_Thesis_Pedro_Porfirio_restructured_v2.2.docx"
OUT = SRC.with_name("Master_Thesis_Pedro_Porfirio_restructured_v2.3.docx")

NEW_HEADING = "5.4  Multi-source integration: layering pipeline-flow data on top of EIA aggregates"

PARA_1 = (
    "The previous two examples both worked from a single authoritative data source: EIA monthly aggregates "
    "at the PADD or USA level. In commercial practice this is not how the workflow ends. Once a desk has a "
    "working monthly balance, the next economically valuable refinement is finer-resolution data on individual "
    "physical assets — daily inventory at the Cushing tank farm, daily flow on the Permian outbound pipelines, "
    "vessel positions in the Houston Ship Channel. The dominant commercial source for crude pipeline flows in "
    "North America is Wood Mackenzie's Genscape service, which reports daily flow rates on named pipelines via "
    "electromagnetic sensors mounted on the lines themselves. Genscape publishes named-pipeline coverage of "
    "approximately ninety-one percent of capacity connected to Cushing and ninety percent of capacity crossing "
    "the United States–Canada border, including Enbridge Mainline, the Keystone system, the Permian outbound "
    "complex (Gray Oak, Cactus II, BridgeTex, Midland-to-ECHO), and the Cushing connectors (Seaway, Spearhead, "
    "Pony Express)."
)

PARA_2 = (
    "Adding Genscape data to a working EIA spreadsheet looks straightforward on first inspection — append "
    "columns, point the formulas at the new cells, recompute the balance — but the integration breaks on five "
    "distinct mismatches that the spreadsheet has no schema-level way to resolve. The frequencies disagree: "
    "Genscape reports daily, EIA monthly, and a Genscape pipeline that is down for four days mid-month leaves "
    "no trace in its monthly mean, so the analyst must choose between losing the operational signal and adding "
    "a parallel daily sheet whose totals must still reconcile to the monthly aggregates upstream. The geographies "
    "disagree: Genscape names a pipeline (DAPL, Keystone) while EIA names a PADD pair (movements from PADD 2 to "
    "PADD 3), and a single EIA inter-PADD figure aggregates over an undisclosed set of pipelines, some of which "
    "Genscape names individually and some of which it does not. Double-counting becomes a live risk on every new "
    "column: if DAPL flows are appended as a separate series and the EIA PADD-2-to-PADD-3 movement series already "
    "incorporates them, every barrel on DAPL is now in the balance twice unless the analyst manually subtracts "
    "it from one side. Coverage is incomplete by design: the nine percent of Cushing-connected capacity that "
    "Genscape does not monitor still shows up in the EIA aggregate, so a residual-pipeline column has to be "
    "maintained by hand each month, recomputed from the difference between the EIA inter-PADD movement and the "
    "sum of Genscape-named flows for the same corridor. Finally, the integration is not stable: every time "
    "Genscape adds coverage of a new pipeline — Seminole-Red, EPIC Crude, the latest reversed segment — the "
    "analyst must add a column, decide which existing aggregate it should be subtracted from, and update the "
    "residual definition. None of these decisions are encoded; they live in the analyst's head and in the "
    "workbook's commented cells, and they evolve as personnel rotate. The cumulative effect is that adding a "
    "finer-resolution data source to the spreadsheet does not deepen the balance; it makes the balance more "
    "fragile, because the same physical crude is now traced through two reconciliation paths with no schema-level "
    "guarantee that the paths agree."
)

PARA_3 = (
    "Table 5.3 illustrates the difference on a representative month for PADD 2, the Midwest PADD whose stock "
    "identity is anchored on Cushing and whose pipeline complex is the most heavily Genscape-monitored in the "
    "country. The figures are illustrative of the order of magnitude observed in commercial pipeline-flow "
    "reports; precise values depend on the subscription product and the month."
)

# Table 5.3 caption — bold "Table 5.3. " + italic explanatory text
TABLE_CAPTION_BOLD = "Table 5.3. "
TABLE_CAPTION_ITALIC = (
    "Layering Genscape pipeline-flow data on top of EIA monthly aggregates for PADD 2, representative month. "
    "The spreadsheet column accumulates one new cell per named pipeline, a residual cell for un-named Genscape "
    "coverage gaps, and a manual subtraction from the EIA inter-PADD column to prevent double-counting. The "
    "framework column TS-binds each Genscape flow to a named pipeline node and the EIA aggregate to the "
    "PADD-view node; the aggregation view ties them and flags divergence automatically."
)

TABLE_ROWS = [
    ("Item", "Spreadsheet workflow", "Framework"),
    ("EIA PADD-2 stock change (monthly)",
     "Single cell in the PADD-2 sheet",
     "padd2_view.inventory TS-bound to MCRSTP21"),
    ("EIA PADD-3-to-PADD-2 movements (monthly)",
     "Single cell; later partly back-subtracted",
     "padd3.outflow → padd2.inflow TS-bound to inter-PADD series"),
    ("Genscape DAPL flow (daily, monthly mean)",
     "New column \"DAPL kbd\"; subtract from EIA movement",
     "dapl.outflow → cushing.inflow TS-bound to Genscape series"),
    ("Genscape Keystone flow (daily, monthly mean)",
     "New column \"Keystone kbd\"; subtract from EIA movement",
     "keystone.outflow → patoka.inflow TS-bound to Genscape series"),
    ("Genscape Seaway-reversed flow",
     "New column; sign convention has to be tracked manually",
     "Reverse-direction edge with its own flow variable; sign is structural"),
    ("Un-named pipelines (≈9 % of capacity)",
     "\"Residual pipelines\" column, recomputed manually each month",
     "Emerges as the difference on the PADD-view node automatically"),
    ("Cross-check that ΣGenscape ≈ EIA inter-PADD",
     "Eyeballed; absent unless the analyst writes a hand-coded formula",
     "v_aggregation_consistency flags divergence in resolved output"),
    ("New pipeline coverage (e.g. Seminole-Red goes live)",
     "Add column; rewrite residual definition; rebuild subtractions",
     "Add a node + TS binding; aggregation view picks it up"),
]

PARA_4 = (
    "The structural point is the same as in the previous two examples, but the failure mode is different. In "
    "§5.2 and §5.3 the spreadsheet's residual cell silently absorbed a meaningful physical signal that the "
    "framework would have surfaced. Here the spreadsheet does not absorb the signal — it surfaces all of it — "
    "but each new signal compounds the integration burden, because the spreadsheet's only mechanism for adding "
    "finer-resolution data is to add columns and patch the formulas downstream by hand. The framework treats "
    "Genscape and EIA as observations on different nodes of the same graph: EIA aggregates bind to PADD-view "
    "nodes, Genscape flows bind to pipeline nodes, the aggregation view ties them through the partition formula "
    "on the PADD-view, and the schema-level CHECK on single attribution makes the double-counting failure mode "
    "unrepresentable. The cost of adding a new Genscape-covered pipeline is not a workbook rewrite; it is a new "
    "TS binding on an existing node. The orthogonality between the physical layer and the data layer — "
    "Principle 1 — is what makes this true. The physical asset graph already contains every pipeline of interest "
    "as a node; subscribing to a finer-resolution data source for a subset of those pipelines does not require "
    "any structural change to the graph. It changes the data layer only."
)


def find_heading(d, text_starts_with):
    """Return index of first paragraph whose text starts with the given prefix
    AND whose style is Heading 2 (avoids matching the TOC entries that share
    the same leading text)."""
    for i, p in enumerate(d.paragraphs):
        style = p.style.name if p.style else ""
        if "Heading 2" not in style:
            continue
        if p.text.strip().startswith(text_starts_with):
            return i
    raise LookupError(f"No Heading-2 paragraph starts with: {text_starts_with!r}")


def make_paragraph(text, style_name=None):
    """Build a w:p element with a single run carrying `text`."""
    p = OxmlElement("w:p")
    if style_name:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_name)
        pPr.append(pStyle)
        p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def make_caption_paragraph(bold_text, italic_text):
    """Build the caption paragraph with bold label + italic descriptive text."""
    p = OxmlElement("w:p")
    # Bold run
    r1 = OxmlElement("w:r")
    rPr1 = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    rPr1.append(b)
    r1.append(rPr1)
    t1 = OxmlElement("w:t")
    t1.text = bold_text
    t1.set(qn("xml:space"), "preserve")
    r1.append(t1)
    p.append(r1)
    # Italic run
    r2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    it = OxmlElement("w:i")
    rPr2.append(it)
    r2.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.text = italic_text
    t2.set(qn("xml:space"), "preserve")
    r2.append(t2)
    p.append(r2)
    return p


def make_table(rows, header_shade="DDDDDD"):
    """Build a w:tbl element. First row is the header (shaded + bold)."""
    tbl = OxmlElement("w:tbl")

    # tblPr
    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "999999")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblLook = OxmlElement("w:tblLook")
    tblLook.set(qn("w:val"), "04A0")
    tblPr.append(tblLook)
    tbl.append(tblPr)

    # tblGrid — 3 columns
    tblGrid = OxmlElement("w:tblGrid")
    for w in ("2400", "3300", "3300"):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), w)
        tblGrid.append(gc)
    tbl.append(tblGrid)

    for row_idx, row in enumerate(rows):
        tr = OxmlElement("w:tr")
        for cell_idx, cell_text in enumerate(row):
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), ("2400", "3300", "3300")[cell_idx])
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            if row_idx == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), header_shade)
                tcPr.append(shd)
            tc.append(tcPr)

            p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            sp = OxmlElement("w:spacing")
            sp.set(qn("w:before"), "40")
            sp.set(qn("w:after"), "40")
            pPr.append(sp)
            p.append(pPr)
            r = OxmlElement("w:r")
            if row_idx == 0:
                rPr = OxmlElement("w:rPr")
                bo = OxmlElement("w:b")
                rPr.append(bo)
                r.append(rPr)
            t = OxmlElement("w:t")
            t.text = cell_text
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl


def renumber_heading(p_obj, new_text):
    """Replace the text of a paragraph that has style Heading 2."""
    # Clear all <w:t> in all runs, set first one to new_text, drop rest.
    runs = p_obj.runs
    if not runs:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def main():
    d = docx.Document(str(SRC))

    # Locate the current §5.4 heading — the insertion target. The replacement
    # operation will move new elements to sit BEFORE this paragraph.
    idx_54 = find_heading(d, "5.4")
    idx_55 = find_heading(d, "5.5")
    p_54 = d.paragraphs[idx_54]
    p_55 = d.paragraphs[idx_55]
    target_elem = p_54._element

    # Build all new elements (heading + 3 prose + caption + table + 1 prose) in
    # the order they should appear, then insert each before target.
    new_elements = [
        make_paragraph(NEW_HEADING, style_name="Heading2"),
        make_paragraph(PARA_1),
        make_paragraph(PARA_2),
        make_paragraph(PARA_3),
        make_caption_paragraph(TABLE_CAPTION_BOLD, TABLE_CAPTION_ITALIC),
        make_table(TABLE_ROWS),
        make_paragraph(""),  # spacer after table
        make_paragraph(PARA_4),
    ]
    for elem in new_elements:
        target_elem.addprevious(elem)

    # Renumber the two displaced headings.
    renumber_heading(p_54, "5.5  The same pattern on other named events")
    renumber_heading(p_55, "5.6  What the comparison establishes")

    # Update the previous Table 5.3 caption to Table 5.4 (the table that was in
    # §5.4 now lives in §5.5; its numbering shifts up by one).
    # Also update any in-prose reference of "Table 5.3" in the §5.5 body to
    # "Table 5.4". Search in the paragraphs between §5.5 and §5.6.
    for i in range(idx_54, idx_55 + 12):
        if i >= len(d.paragraphs):
            break
        for run in d.paragraphs[i].runs:
            if run.text and "Table 5.3" in run.text:
                run.text = run.text.replace("Table 5.3", "Table 5.4")

    d.save(str(OUT))
    print(f"Spliced new §5.4 into {OUT.name}")
    print(f"Renumbered displaced sections to §5.5 and §5.6")


if __name__ == "__main__":
    main()
