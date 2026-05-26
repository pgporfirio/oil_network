"""One-shot: add bold "Principle N." callouts to the first body paragraph
of §4.1-§4.6 in Master_Thesis_Pedro_Porfirio_restructured_v2.0.docx.

Headings and TOC are left unchanged. The callout sits inline at the start of
each section's first paragraph, so cross-references and structure are
unaffected. Writes to v2.1.docx alongside the source.
"""
from __future__ import annotations
from pathlib import Path

import docx
from docx.oxml.ns import qn

SRC = Path(__file__).resolve().parent.parent / "outputs" / "docs" / "Master_Thesis_Pedro_Porfirio_restructured_v2.0.docx"
OUT = SRC.with_name("Master_Thesis_Pedro_Porfirio_restructured_v2.1.docx")


def main():
    d = docx.Document(str(SRC))
    paragraphs = d.paragraphs

    added = 0
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        style = p.style.name if p.style else ""
        if "Heading 2" in style:
            heading = p.text.strip()
            # Match "4.X " where 1 <= X <= 6
            try:
                num, _rest = heading.split(" ", 1)
                major, minor = num.split(".")
                minor = int(minor)
            except (ValueError, IndexError):
                i += 1
                continue
            if major == "4" and 1 <= minor <= 6:
                # First non-empty body paragraph after the heading
                for j in range(i + 1, len(paragraphs)):
                    target = paragraphs[j]
                    if target.text.strip():
                        # Add a new bold run at the END, then move it to be the
                        # first <w:r> child in the paragraph.
                        callout = target.add_run(f"Principle {minor}. ")
                        callout.bold = True
                        r_elem = callout._element
                        p_elem = target._element
                        first_r = None
                        for child in p_elem:
                            if child.tag == qn("w:r"):
                                first_r = child
                                break
                        if first_r is not None and first_r is not r_elem:
                            p_elem.remove(r_elem)
                            first_r.addprevious(r_elem)
                        added += 1
                        break
        i += 1

    d.save(str(OUT))
    print(f"Added 'Principle N.' callouts to {added} sections")
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
