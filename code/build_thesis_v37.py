"""Build Master_Thesis_Pedro_Porfirio_v37.docx from v36.

Changes vs v36:
  (1) Replace Figure 3.1 (paragraph 99) with the correct physical
      supply-chain schematic. The v36 docx accidentally embedded the
      Asset / Node / Graph / Variable / Scenario entity-relationships
      diagram (which still appears, correctly, at the end of Chapter 4
      as Figure 1). Fix: swap /word/media/image1.png in the docx zip
      with outputs/docs/fig_3_1_physical_supply_chain.png. The new PNG
      is rendered at the same 1.654 aspect ratio as the slot's stored
      wp:extent, so no document.xml edit is required.

  (2) Update Section 10.3 (Future work) to reflect the post-v36 state
      where the U.S. crude grade registry is now seeded in the
      database (19 grades + a parent-child hierarchy + a
      v_commodity_ancestors view) while the resolver continues to
      treat commodity = crude as a single dimension. Keeps the prose
      honest about what is in place today versus what remains for
      future per-grade decomposition work.

Word count: ~+30 words (one inserted sentence in 10.3).
"""
from __future__ import annotations

import os
import shutil
import sys
import zipfile

from docx import Document

CLEAN_DIR = os.path.normpath(os.path.join(os.path.dirname(__file__), '..'))
DOCS_DIR = os.path.join(CLEAN_DIR, 'outputs', 'docs')
SRC = os.path.join(DOCS_DIR, 'Master_Thesis_Pedro_Porfirio_v36.docx')
DST = os.path.join(DOCS_DIR, 'Master_Thesis_Pedro_Porfirio_v37.docx')
FIG_3_1 = os.path.join(DOCS_DIR, 'fig_3_1_physical_supply_chain.png')


def _add_run_with_style(paragraph, text: str):
    """Append a run that matches the formatting of the paragraph's last run."""
    run = paragraph.add_run(text)
    if paragraph.runs and len(paragraph.runs) > 1:
        ref = paragraph.runs[0]
        run.font.name = ref.font.name
        run.font.size = ref.font.size
        run.bold = ref.bold
        run.italic = ref.italic
    return run


def edit_future_work(doc: Document) -> None:
    """Insert the grade-registry sentence into the Section 10.3 paragraph.

    The existing paragraph already says: 'The schema also admits per-grade
    decomposition by treating commodity as a further dimension on variables,
    a weekly companion scenario built on EIA's WPSR, ...'. We insert a
    follow-on clause noting that the registry is now in place but the
    resolver has not yet been extended to propagate grades.
    """
    target_anchor = "The schema also admits per-grade decomposition"
    inserted_sentence = (
        " A first step towards this extension is now in place: the schema "
        "carries a seeded crude-grade registry with nineteen named U.S. grades "
        "(WTI and its delivery points, Bakken, Eagle Ford light and condensate, "
        "LLS, Niobrara, Oklahoma sweet, the Gulf medium-sour family, Alaska "
        "North Slope, and the California heavy family), together with a "
        "parent-child hierarchy and a recursive ancestor view. The resolver, "
        "however, continues to treat commodity = crude as a single dimension; "
        "extending it to propagate per-grade balances is the next concrete "
        "deliverable beyond the thesis."
    )
    for para in doc.paragraphs:
        if target_anchor in para.text:
            # Append the sentence as a new run on the same paragraph so the
            # paragraph stays as one flowing piece of prose.
            _add_run_with_style(para, inserted_sentence)
            return
    raise RuntimeError("Could not find Section 10.3 anchor for grade-registry sentence")


def replace_image_in_zip(docx_path: str, member: str, replacement_path: str) -> None:
    """Replace one member of the docx zip in-place."""
    import tempfile
    with open(replacement_path, 'rb') as f:
        new_bytes = f.read()
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=os.path.dirname(docx_path))
    os.close(tmp_fd)
    try:
        with zipfile.ZipFile(docx_path, 'r') as zin, \
             zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            replaced = False
            for item in zin.infolist():
                if item.filename == member:
                    zout.writestr(item, new_bytes)
                    replaced = True
                else:
                    zout.writestr(item, zin.read(item.filename))
            if not replaced:
                raise RuntimeError(f"Member {member} not found in docx")
        shutil.move(tmp_path, docx_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def main() -> int:
    if not os.path.exists(SRC):
        print(f"ERROR: source not found: {SRC}", file=sys.stderr)
        return 1
    if not os.path.exists(FIG_3_1):
        print(f"ERROR: figure not found: {FIG_3_1}", file=sys.stderr)
        print("Run code/build_fig_3_1.py first.", file=sys.stderr)
        return 1

    # 1. Copy v36 -> v37
    shutil.copyfile(SRC, DST)
    print(f"[1/3] copied {os.path.basename(SRC)} -> {os.path.basename(DST)}")

    # 2. Edit Section 10.3 via python-docx
    doc = Document(DST)
    edit_future_work(doc)
    doc.save(DST)
    print("[2/3] inserted grade-registry sentence in Section 10.3")

    # 3. Replace embedded image1.png with the new supply-chain schematic
    replace_image_in_zip(DST, 'word/media/image1.png', FIG_3_1)
    print(f"[3/3] replaced word/media/image1.png with {os.path.basename(FIG_3_1)}")

    print(f"\nDONE: {DST}")
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
